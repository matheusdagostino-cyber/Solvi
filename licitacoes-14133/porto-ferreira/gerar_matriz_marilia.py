#!/usr/bin/env python3
"""Gera a Matriz de Argumentos (formato Marília) em DOCX.

Lê os blocos AN_bloco{1..6}.md e monta a tabela conforme CLAUDE.md:
| Nº | Tema | Problema | Argumento | Fundamento | Esclarecimentos? | Aplicação |

Pontos [RESERVA] vão em seção segregada ao final.
Flags [VIT] e [VALIDAÇÃO TÉCNICA] preservados.
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

DOCS = Path(__file__).parent / "docs"
BLOCOS = sorted(DOCS.glob("AN_bloco*.md"))

TEMA_MAP = {
    "T1": "MODELAGEM_ECONOMICA",
    "T2": "OBJETO_ESCOPO",
    "T3": "OBJETO_ESCOPO",
    "T4": "MODELAGEM_ECONOMICA",
    "T5": "MODELAGEM_ECONOMICA",
    "T6": "MODELAGEM_ECONOMICA",
    "T7": "LICENCIAMENTO_AMBIENTAL",
    "T8": "OBJETO_ESCOPO",
    "T9": "OBJETO_ESCOPO",
    "T10": "RISCOS_GARANTIAS",
    "T11": "REGULACAO_FISCALIZACAO",
    "T12": "REMUNERACAO_REAJUSTE",
    "T13": "HABILITACAO",
    "T14": "HABILITACAO",
    "T15": "QUESTOES_FORMAIS",
    "T16": "REGULACAO_FISCALIZACAO",
    "T17": "QUESTOES_FORMAIS",
    "T18": "QUESTOES_FORMAIS",
}

TEMA_LABEL = {
    "OBJETO_ESCOPO": "Objeto e escopo",
    "HABILITACAO": "Habilitação",
    "JULGAMENTO_PROPOSTA": "Julgamento",
    "MODELAGEM_ECONOMICA": "Modelagem econômica",
    "RISCOS_GARANTIAS": "Riscos e garantias",
    "REMUNERACAO_REAJUSTE": "Remuneração e reajuste",
    "REGULACAO_FISCALIZACAO": "Regulação e fiscalização",
    "LICENCIAMENTO_AMBIENTAL": "Licenciamento ambiental",
    "QUESTOES_FORMAIS": "Questões formais",
}


def parse_bloco(path):
    text = path.read_text(encoding="utf-8")
    sections = re.split(r"(?=^## )", text, flags=re.MULTILINE)
    items = []
    for sec in sections:
        sec = sec.strip()
        if not sec.startswith("## "):
            continue
        header_line = sec.split("\n", 1)[0]
        # T-sections
        m_t = re.match(r"## (T\d+)\s*[—–-]\s*(.+)", header_line)
        # R-sections (reserves)
        m_r = re.match(r"## (R\d+)\s*\(([^)]+)\)\s*[—–-]\s*(.+)", header_line)
        if m_r:
            rid = m_r.group(1)
            cids = m_r.group(2)
            title = m_r.group(3).strip()
            body = sec.split("\n", 1)[1] if "\n" in sec else ""
            items.append({
                "type": "reserve",
                "id": rid,
                "cids": cids,
                "title": title,
                "body": body,
            })
        elif m_t:
            tid = m_t.group(1)
            full_title = m_t.group(2).strip()
            # extract short title (up to first colon or parenthesis detail)
            short = re.split(r"[:(]", full_title)[0].strip().rstrip(" —–-")
            body = sec.split("\n", 1)[1] if "\n" in sec else ""
            items.append({
                "type": "thesis",
                "id": tid,
                "title": full_title,
                "short_title": short,
                "body": body,
            })
    return items


def extract_section(body, heading):
    pattern = rf"^### {re.escape(heading)}\s*\n(.*?)(?=^### |\Z)"
    m = re.search(pattern, body, re.MULTILINE | re.DOTALL)
    if m:
        return m.group(1).strip()
    return ""


def extract_roteamento_field(roteamento_text, field):
    pattern = rf"- {re.escape(field)}:\s*(.+?)(?:\n|$)"
    m = re.search(pattern, roteamento_text)
    if m:
        return m.group(1).strip()
    return ""


def extract_flags(body):
    flags_line = ""
    for line in body.split("\n")[:5]:
        if line.startswith("**Flags:**"):
            flags_line = line
            break
    return flags_line


def set_cell_text(cell, text, font_name="Times New Roman", font_size=10,
                  bold=False, alignment=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    # line spacing
    p.paragraph_format.line_spacing = Pt(13)


def set_cell_shading(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = tcPr.find(qn("w:shd"))
    if shading is None:
        shading = tc.makeelement(qn("w:shd"), {})
        tcPr.append(shading)
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")


def add_multiline_cell(cell, text, font_name="Times New Roman", font_size=10):
    cell.text = ""
    lines = text.split("\n")
    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.line_spacing = Pt(13)

        # handle bold markers and [VIT]/[VALIDAÇÃO TÉCNICA] highlights
        parts = re.split(r"(\*\*[^*]+\*\*|\[VIT\]|\[VALIDAÇÃO TÉCNICA[^\]]*\])", line)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part in ("[VIT]",) or part.startswith("[VALIDAÇÃO TÉCNICA"):
                run = p.add_run(part)
                run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                run.bold = True
            else:
                run = p.add_run(part)
            run.font.name = font_name
            run.font.size = Pt(font_size)


def build_argumento_text(body):
    arg = extract_section(body, "Argumento")
    contra = extract_section(body, "Contra-argumentos")
    combined = arg
    if contra:
        combined += "\n\n— Contra-argumentos antecipados —\n\n" + contra
    # trim excessive length for the table
    if len(combined) > 8000:
        combined = combined[:8000] + "\n[...texto completo no arquivo AN_bloco correspondente]"
    return combined


def build_fundamento_text(body, jusia_extra=""):
    fund = extract_section(body, "Fundamento")
    doutrina = extract_section(body, "Doutrina")
    parts = [fund]
    if doutrina and doutrina != "—":
        parts.append("Doutrina: " + doutrina)
    if jusia_extra:
        parts.append(jusia_extra)
    return "\n".join(p for p in parts if p)


FORO_MAP = {
    "ADMINISTRACAO": "Administração",
    "TCE": "TCE",
    "JUDICIARIO": "Judiciário",
}


def parse_annex_by_theme(text):
    """Split an annex markdown file into a dict keyed by T-number."""
    sections = re.split(r"(?=^## T\d+)", text, flags=re.MULTILINE)
    result = {}
    for sec in sections:
        m = re.match(r"^## (T\d+)", sec)
        if m:
            tid = m.group(1)
            result[tid] = sec.strip()
    return result


def summarize_annex_for_fundamento(annex_section, max_chars=2000):
    """Extract key decisions from an annex section for the Fundamento column."""
    if not annex_section:
        return ""
    lines = []
    for line in annex_section.split("\n"):
        line_s = line.strip()
        if line_s.startswith("- ") and ("[VIT]" in line_s or "Acórdão" in line_s
                                        or "Entendimento" in line_s
                                        or "Súmula" in line_s):
            entry = line_s[2:].strip()
            dash = entry.find(" — ", 0, 120)
            if dash > 0:
                entry = entry[:dash]
            lines.append("• " + entry)
    if not lines:
        return ""
    combined = "\n".join(lines)
    if len(combined) > max_chars:
        combined = combined[:max_chars] + "\n[...ver anexo completo]"
    return combined


def main():
    # Parse all blocks
    all_theses = []
    all_reserves = []

    for bloco in BLOCOS:
        items = parse_bloco(bloco)
        for item in items:
            if item["type"] == "thesis":
                item["_source"] = bloco.name
                all_theses.append(item)
            elif item["type"] == "reserve":
                item["_source"] = bloco.name
                all_reserves.append(item)

    # Load JusIA annex if available
    jusia_path = DOCS / "JusIA_anexo.md"
    jusia_text = jusia_path.read_text(encoding="utf-8") if jusia_path.exists() else ""
    jusia_by_t = parse_annex_by_theme(jusia_text) if jusia_text else {}

    # Load Lei na Mão annex if available
    lnm_path = DOCS / "TCE_LeiNaMao_anexo.md"
    lnm_text = lnm_path.read_text(encoding="utf-8") if lnm_path.exists() else ""
    lnm_by_t = parse_annex_by_theme(lnm_text) if lnm_text else {}

    print(f"Parsed {len(all_theses)} theses, {len(all_reserves)} reserves")
    if jusia_by_t:
        print(f"JusIA annex: {len(jusia_by_t)} themes loaded")
    if lnm_by_t:
        print(f"Lei na Mão annex: {len(lnm_by_t)} themes loaded")

    # Create document
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)  # landscape A4
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("PE 30/2026 (REABERTURA) — Porto Ferreira/SP")
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.bold = True

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Matriz de Argumentos — Fase 2")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = True

    # Header info
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.LEFT
    info_text = (
        "[USO INTERNO]\n"
        "Órgão: Prefeitura Municipal de Porto Ferreira/SP\n"
        "Modalidade: Pregão Eletrônico n.º 30/2026 (Reabertura)\n"
        "Objeto: Contratação de serviços de coleta, transporte e "
        "destinação final de RSU, com operação de aterro sanitário "
        "municipal e fornecimento de contêineres\n"
        "Regime: Lei n.º 14.133/2021 — Serviço comum contínuo\n"
        "Data de referência: agosto/2026\n"
        "Pipeline: Fase 1 (E1-E4 → C: 300 pontos) → "
        "Fase 2 (AN+R: 18 macro-questões + 7 reservas)\n"
        "Jurisprudência: Lei na Mão (TCU/TCE-SP) + Jus IA (STJ/TJs). "
        "Citações de busca automatizada marcadas [VIT]."
    )
    run = info.add_run(info_text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    info.paragraph_format.space_after = Pt(12)

    doc.add_paragraph()  # spacer

    # Main table
    cols = ["Nº", "Tema", "Problema", "Argumento", "Fundamento",
            "Escl.?", "Aplicação"]
    col_widths = [Cm(1.2), Cm(2.3), Cm(3.3), Cm(10.0), Cm(6.7),
                  Cm(1.2), Cm(2.0)]

    table = doc.add_table(rows=1, cols=len(cols))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, col_name in enumerate(cols):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, col_name, bold=True, font_size=9,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, "1F4E79")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Set column widths
    for i, width in enumerate(col_widths):
        for row in table.rows:
            row.cells[i].width = width

    # Add thesis rows
    for idx, thesis in enumerate(all_theses, 1):
        tid = thesis["id"]
        body = thesis["body"]
        row = table.add_row()

        # Nº
        set_cell_text(row.cells[0], tid, font_size=9,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)

        # Tema
        tema_key = TEMA_MAP.get(tid, "")
        tema_label = TEMA_LABEL.get(tema_key, tema_key)
        set_cell_text(row.cells[1], tema_label, font_size=9)

        # Problema (short title + flags)
        flags = extract_flags(body)
        problema = thesis["short_title"]
        if "[VALIDAÇÃO TÉCNICA]" in flags or "[VALIDAÇÃO TÉCNICA]" in body[:500]:
            problema += "\n[VALIDAÇÃO TÉCNICA]"
        add_multiline_cell(row.cells[2], problema, font_size=9)

        # Argumento (full text with contra-arguments)
        arg_text = build_argumento_text(body)
        add_multiline_cell(row.cells[3], arg_text, font_size=8)

        # Fundamento (merge AN block + annex jurisprudence)
        annex_parts = []
        lnm_sec = lnm_by_t.get(tid, "")
        if lnm_sec:
            summary = summarize_annex_for_fundamento(lnm_sec, 1500)
            if summary:
                annex_parts.append("— Jurisprudência TCU/TCE (Lei na Mão) —\n" + summary)
        jusia_sec = jusia_by_t.get(tid, "")
        if jusia_sec:
            summary = summarize_annex_for_fundamento(jusia_sec, 1500)
            if summary:
                annex_parts.append("— Jurisprudência STJ/TJ + doutrina (Jus IA) —\n" + summary)
        jusia_extra = "\n\n".join(annex_parts)
        fund_text = build_fundamento_text(body, jusia_extra)
        add_multiline_cell(row.cells[4], fund_text, font_size=8)

        # Esclarecimentos?
        roteamento = extract_section(body, "Roteamento")
        escl = extract_roteamento_field(roteamento, "esclarecimento")
        escl_val = "Sim" if escl.lower().startswith("sim") else "Não"
        set_cell_text(row.cells[5], escl_val, font_size=9,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)

        # Aplicação
        foro = extract_roteamento_field(roteamento, "foro_sugerido")
        canal = extract_roteamento_field(roteamento, "canal_sugerido")
        foro_clean = foro.split("(")[0].strip().split(",")[0].strip()
        foro_label = FORO_MAP.get(foro_clean, foro_clean)
        aplicacao = foro_label
        if canal:
            canal_short = canal.split("—")[0].split("(")[0].strip()
            aplicacao = f"{canal_short}\n({foro_label})"
        set_cell_text(row.cells[6], aplicacao, font_size=9,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)

        # Alternate row shading
        if idx % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, "F2F2F2")

    # RESERVA section
    if all_reserves:
        doc.add_paragraph()
        sep = doc.add_paragraph()
        sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sep.add_run("[RESERVA — NÃO PROTOCOLAR]")
        run.font.name = "Times New Roman"
        run.font.size = Pt(13)
        run.bold = True
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

        note = doc.add_paragraph()
        run = note.add_run(
            "Pontos que funcionam como vantagem competitiva para o grupo. "
            "A decisão de levantar ou reservar é exclusiva do advogado."
        )
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.italic = True

        for res in all_reserves:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            run = p.add_run(f"{res['id']} ({res['cids']}) — {res['title']}")
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            run.bold = True

            body_p = doc.add_paragraph()
            # Extract the analysis text
            analysis = res["body"].strip()
            if len(analysis) > 3000:
                analysis = analysis[:3000] + "\n[...ver AN_bloco6.md]"
            run = body_p.add_run(analysis)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            body_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Footer
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run(
        "Documento gerado automaticamente pelo pipeline Solvi (Fase 2 AN+R). "
        "Jurisprudência [VIT] pendente de verificação no inteiro teor. "
        "Pontos [VALIDAÇÃO TÉCNICA] dependem de confirmação de engenharia."
    )
    run.font.name = "Times New Roman"
    run.font.size = Pt(8)
    run.italic = True

    # Save
    out_path = DOCS / "MATRIZ_ARGUMENTOS_PE30_2026_PortoFerreira.docx"
    doc.save(str(out_path))
    print(f"\nDocumento salvo: {out_path}")
    print(f"  Teses: {len(all_theses)}")
    print(f"  Reservas: {len(all_reserves)}")


if __name__ == "__main__":
    main()
